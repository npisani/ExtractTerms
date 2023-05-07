#!/usr/bin/env python3

import os
import nltk
import docx
from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from lxml import etree
import nltk
nltk.download('punkt')

def extract_words(document):
    """Extracts words from a Word document."""
    words = []

    # Extract words from paragraphs
    for paragraph in document.paragraphs:
        words.extend(custom_tokenize(paragraph.text))

    # Extract words from table cells
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    words.extend(custom_tokenize(paragraph.text))

    # Extract text from text boxes
    text_box_text = extract_text_from_text_boxes(document)
    words.extend(custom_tokenize(text_box_text))

    return words


def extract_text_from_text_boxes(document):
    text = ""
    namespaces = {
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    root = etree.fromstring(document.element.xml)
    for shape in root.xpath("//wps:txbxContent", namespaces=namespaces):
        for paragraph in shape.xpath(".//w:p", namespaces=namespaces):
            paragraph_obj = Paragraph(paragraph, document)
            text += paragraph_obj.text + " "
    return text

def custom_tokenize(text):
    """Tokenize text considering brackets, periods, and other special characters as delimiters."""
    tokens = nltk.word_tokenize(text)
    expanded_tokens = []
    for token in tokens:
        if token.startswith('[') and token.endswith(']'):
            token = token[1:-1]  # Remove square brackets
            expanded_tokens.extend(token.split('.'))
        else:
            expanded_tokens.append(token)

    # Replace placeholders with the expanded tokens and remove preceding special characters
    cleaned_tokens = []
    for i, token in enumerate(expanded_tokens):
        # Remove preceding special characters
        token = re.sub(r'^[-.,]+', '', token)
        cleaned_tokens.append(token)

    return [token for token in cleaned_tokens if token]  # Move this return statement inside the function

def get_unique_terms(words):
    """Extracts unique terms from a list of words."""
    unique_terms = set()
    for word in words:
        # Updated regex pattern to capture the mentioned terms and others
        if re.match(r'\b(?:[A-Z]{1,2}\d{1,2}[A-Z]{0,2}-?\d{0,5}|[A-Z]{2,}\d{1,2}(?:\.\d{2})?|[A-Z]{1}\d{1}[A-Z]{1}\d{1}|(?=.*[A-Za-z])[A-Za-z0-9._\'\[\]-]+(?<!\d)$)\b', word) and len(word) > 1:
            unique_terms.add(word.lower())  # convert to lowercase
    return unique_terms

def save_sensitive_terms(terms, input_file_name):
    """Saves a set of sensitive terms to a file."""
    output_directory = 'Output Files'
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    sensitive_terms_file = os.path.join(output_directory, f"{input_file_name}_Sensitive Terms List.txt")
    with open(sensitive_terms_file, 'w', encoding='utf-8') as f:
        for term in sorted(terms):
            f.write(f"{term}\n")

def read_existing_sensitive_terms(input_file):
    """Reads sensitive terms from an existing 'Sensitive Terms List.txt' file in the output directory."""
    sensitive_terms = set()
    input_file_name = os.path.splitext(os.path.basename(input_file))[0]
    output_directory = 'Output Files'
    sensitive_terms_file = os.path.join(output_directory, f"{input_file_name}_Sensitive Terms List.txt")

    if os.path.exists(sensitive_terms_file):
        with open(sensitive_terms_file, 'r', encoding='utf-8') as f:
            for line in f:
                sensitive_terms.add(line.strip().lower())  # Convert the term to lowercase before adding
    return sensitive_terms

def show_sensitive_terms_window(unique_terms, existing_sensitive_terms, input_file):
    """Displays a window with a list of unique terms with checkboxes and a save/cancel button."""
    sensitive_terms = set(existing_sensitive_terms)

    def save():
        nonlocal sensitive_terms
        for i, term in enumerate(sorted(unique_terms)):
            if var_list[i].get() == 1:
                sensitive_terms.add(term)
        input_file_name = os.path.splitext(os.path.basename(input_file))[0]
        save_sensitive_terms(sensitive_terms, input_file_name)
        window.destroy()

    def cancel():
        window.destroy()

    if not unique_terms:
        print("No unique terms found.")
        return

    window = Tk()
    window.title("Sensitive Terms")
    window.geometry("500x600")

    frame1 = Frame(window)
    frame1.pack(side=TOP, fill=BOTH, expand=1)

    frame2 = Frame(window)
    frame2.pack(side=BOTTOM, pady=10)

    canvas = Canvas(frame1)
    scrollbar = Scrollbar(frame1, orient=VERTICAL, command=canvas.yview)
    scrollable_frame = Frame(canvas)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Add mouse wheel scrolling
    def on_mouse_wheel(event):
        canvas.yview_scroll(-1 * int((event.delta / 120)), "units")

    canvas.bind_all("<MouseWheel>", on_mouse_wheel)

    canvas.pack(side=LEFT, fill=BOTH, expand=1)
    scrollbar.pack(side=RIGHT, fill=Y)

    var_list = []
    # Modify the checkboxes loop to pre-fill the existing sensitive terms
    for term in sorted(unique_terms):
        var = IntVar()
        var_list.append(var)
        chk = Checkbutton(scrollable_frame, variable=var, onvalue=1, offvalue=0, text=term, anchor=W)
        if term.lower() in existing_sensitive_terms:  # Convert the term to lowercase before checking
            chk.select()
        chk.pack(side=TOP, fill=X)

    save_button = ttk.Button(frame2, text="SAVE", command=save)
    save_button.pack(side=LEFT, padx=10)

    cancel_button = ttk.Button(frame2, text="CANCEL", command=cancel)
    cancel_button.pack(side=LEFT, padx=10)

    window.mainloop()

def main():
    # Prompt user to select Word document
    input_file = filedialog.askopenfilename(title="Select Word Document", filetypes=[("Word Document", "*.docx")])

    # Check if file was selected
    if not input_file:
        print("No file selected.")
    else:
        # Extract unique terms and display in window
        document = Document(input_file)
        words = extract_words(document)
        unique_terms = get_unique_terms(words)
        existing_sensitive_terms = read_existing_sensitive_terms(input_file)
        show_sensitive_terms_window(unique_terms, existing_sensitive_terms, input_file)

if __name__ == "__main__":
    main()
